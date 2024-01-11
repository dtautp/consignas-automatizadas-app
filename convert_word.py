import pandas as pd
import numpy as np
from docx import Document
from docx2pdf import convert
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def excel_to_word(excel_file, ruta_destino):
    # Cargar datos desde el archivo Excel'
    df = pd.read_excel(excel_file)

    # Cambiar todos los datos a tipo string en el DataFrame
    df = df.astype(str)

    # Definir una función para transformar los datos de puntos
    def transformar_valor(valor):
        if '.' in valor:
            parte_decimal = valor.split('.')[1]
            if parte_decimal == '0':
                return valor.split('.')[0]
        return valor

    # Obtener las columnas que terminan en "*pts"
    columnas_pts = [col for col in df.columns if col.endswith('pts')]

    # Aplicar la transformación solo a las columnas que terminan en "*pts"
    df[columnas_pts] = df[columnas_pts].apply(lambda col: col.map(transformar_valor))

    # Plantillas de Word
    word_template_evaluate = 'input/templates/Consigna de actividad (Evaluada).docx'
    word_template_not_evaluate = 'input/templates/Consigna de actividad (No Evaluada).docx'

    # Verificar si existen las plantillas de Word
    templates_exist = all([
        os.path.exists(word_template_evaluate),
        os.path.exists(word_template_not_evaluate)
    ])

    if not templates_exist:
        message = f"No se encontraron una o ambas plantillas de Word"
        return (message)

    # Recorrer filas del DataFrame
    for index, row in df.iterrows():
        # Crear un nuevo documento Word basado en la plantilla de consigna evaluada
        doc = Document(word_template_evaluate)
        # Si no es eveluada eleguir al otra plantilla no evaluada
        if row['Es evaluada?'] == 'No':
            doc = Document(word_template_not_evaluate)
        
        # Obtener valores para el nombre del archivo
        titulo_consigna = row['Título consigna de actividad']
        id_rubrica = row['id_rubrica']

        # Construir 'nombre_archivo' considerando valores nulos en 'id_rubrica'
        nombre_archivo = titulo_consigna
        if id_rubrica != 'nan':
            nombre_archivo = titulo_consigna + '_' + id_rubrica

        # Buscar y reemplazar en párrafos
        for para in doc.paragraphs:
            for key, value in row.items():
                if '{' + key + '}' in para.text:  # Verifica si la llave está en el texto del párrafo
                    lines = value.split('\n')  # Divide el valor en líneas
                    if lines:
                        para.text = para.text.replace(f'{{{key}}}', lines[0])  # Reemplaza la primera línea
                        for line in lines[1:]:
                            # Crear nuevo párrafo en la posición correcta
                            new_para = doc.add_paragraph(style=para.style)
                            para._element.addnext(new_para._element)
                            # Detectar si el texto comienza con "-" y aplicar estilo de viñeta
                            if line.strip().startswith('-'):
                                # Remover el "-" del inicio y los espacios adicionales alrededor
                                new_para.text = line.strip()[1:].strip()
                                if 'Lista viñeta' in doc.styles:
                                    new_para.style = 'Lista viñeta'  # Asegúrate de que el nombre coincida con el estilo que creaste
                                
                            else:
                                new_para.text = line
                                # Asegúrate de restablecer el estilo a 'Normal' o a otro que desees
                                new_para.style = 'Normal'  # Ajusta esto al estilo por defecto que deseas para textos no viñeta
                            para = new_para


        # Buscar y reemplazar en tablas
        for table in doc.tables:
            for row_doc in table.rows:
                for cell in row_doc.cells:
                    for key, value in row.items():
                        cell.text = cell.text.replace(f'{{{key}}}', str(value))
        
        # Eliminar filas con 'nan' en las tablas
        for table in doc.tables:
            rows_to_delete = set()  # Usamos un set para evitar duplicados
            for i, row in enumerate(table.rows):
                text_in_row = [cell.text.strip() for cell in row.cells]
                # Si todas las celdas contienen 'nan' o están vacías, marcar para eliminar
                if all(cell_text == 'nan' or cell_text == '' for cell_text in text_in_row):
                    rows_to_delete.add(i)  # Marcar fila actual
                    if i + 1 < len(table.rows):  # Asegurarse de que la siguiente fila existe
                        rows_to_delete.add(i + 1)  # Marcar la siguiente fila
            
            # Eliminar filas en orden inverso para no alterar los índices
            for i in sorted(rows_to_delete, reverse=True):
                table._tbl.remove(table.rows[i]._tr)

        # Iterar sobre las tablas y establecer el formato del texto solo en la primera fila
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if i == 0:  # Verificar si es la primera fila
                                run.font.name = 'Lato'
                                run.font.size = Pt(14)  # Tamaño de la fuente
                                run.bold = True  # Negrita
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alineación
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            elif i >= 1:
                                run.font.size = Pt(14)
                                run.font.name = 'Lato'
                                if i % 2 == 0:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                else:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                # Aplicar un formato diferente a la primera columna
                                if j == 0:
                                    run.font.size = Pt(14)  # Tamaño de la fuente
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    run.bold = True  # Negrita
                                    run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro

        # Reemplazar 'nan' por una cadena vacía en todo el documento
        # Reemplazo en párrafos
        for para in doc.paragraphs:
            if 'nan' in para.text:
                para.text = para.text.replace('nan', '')
        
        # Reemplazo en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if 'nan' in cell.text:
                        cell.text = cell.text.replace('nan', '')

        # Guardar el documento combinado
        output_file = f'{ruta_destino}/{nombre_archivo}.docx'
        doc.save(output_file)

        # Convertir el documento Word a PDF
        pdf_output_file = f'{ruta_destino}/{nombre_archivo}.pdf'
        convert(output_file, pdf_output_file)

    print("Documentos exportados exitosamente.")
    return ("ok")
